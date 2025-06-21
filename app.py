from flask import Flask, render_template, request, jsonify, send_file
from docx import Document
import os
import PyPDF2
import cohere

app = Flask(__name__)

co = cohere.Client("jXVuo5jfzCLp13QMJfVOU4qnfdutyuacWjR8lMcO")




def generateCoverLetter(name, email, phone, jobTitle, company, jobDesc, resumeText):
    prompt = f"""
    Based on the following resume and job description, write a personalized, formal, and tailored cover letter:

    Name: {name}
    Email: {email}
    Phone: {phone}


    Job Title: {jobTitle}
    Company: {company}
    Job Description:
    {jobDesc}

    The cover letter should highlight relevant experiences and skills from the resume and explain why the applicant is a strong fit.
    """

    response = co.generate(
        model='command-r-plus',  
        prompt=prompt,
        max_tokens=500,
        temperature=0.7
    )

    return response.generations[0].text


@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        try:
            data = request.get_json()

            name = data['name']
            email = data['email']
            phone = data['phone']
            jobTitle = data['jobTitle']
            company = data['company']
            jobDesc = data['jobDesc']

            resumeText = "Sample resume text goes here."  # You can integrate file reading later

            cover_letter = generateCoverLetter(name, email, phone, jobTitle, company, jobDesc, resumeText)
            return jsonify({'cover_letter': cover_letter})

        except Exception as e:
            print("SERVER ERROR:", str(e))
            return jsonify({'error': str(e)}), 500

    return render_template('index.html')



@app.route('/download_docx', methods=['POST'])
def download_docx():
    data = request.get_json()
    cover_letter = data['cover_letter']
    
    doc = Document()
    doc.add_heading('Cover Letter', level=1)
    doc.add_paragraph(cover_letter)
    
    fileName = 'Cover_letter.docx'
    doc.save(fileName)
    
    response = send_file(fileName, as_attachment=True)

    return response


if __name__ == "__main__":
    app.run(debug=True)
